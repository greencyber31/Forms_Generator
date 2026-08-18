import pandas as pd
from docxtpl import DocxTemplate
from docx.shared import Mm
import os
import subprocess
import concurrent.futures
import logging
import shutil
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from pypdf import PdfWriter

EXCEL_FILE     = "CFITF Farmers.xlsx"
TEMPLATE_FILE  = "Application for Crop Insurance 2026 final only.docx"
TRANSMITTAL_TEMPLATE = "Transmittal_Template.docx"
OUTPUT_DIR     = Path("output")
TEMP_DIR       = Path("temp_docs")
LOG_DIR        = Path("logs")

# Set to an integer (e.g. 10) to process only the first N rows for testing.
# Set to None to process ALL rows.
TEST_ROW_LIMIT = None

# Number of parallel workers.  A good starting point is your CPU core count.
# Reduce if you run out of RAM or see LibreOffice crashes.
MAX_WORKERS    = 4

# Characters that are invalid or problematic in filenames (Linux + Windows safe)
_UNSAFE_CHARS = str.maketrans({
    '/': '_', '\\': '_', ':': '_', '*': '_',
    '?': '_', '"': '_', '<': '_', '>': '_', '|': '_',
})
# ───────────────────────────────────────────────────────────────────────────────


def sanitize_key(k: str) -> str:
    """Convert an Excel column header into a valid Jinja2 variable name."""
    s = str(k).strip()
    s = s.replace(' ', '_').replace('.', '').replace('-', '_').replace('/', '_')
    return ''.join(c for c in s if c.isalnum() or c == '_')


def safe_filename(name: str) -> str:
    """Replace all OS-reserved / path-unsafe characters in a filename segment."""
    return str(name).translate(_UNSAFE_CHARS).strip()


def _build_context(row_dict: dict) -> dict:
    """Convert a raw row dictionary into a clean Jinja2 rendering context."""
    context = {}
    for k, v in row_dict.items():
        if pd.isna(v):
            clean_v = ""
        elif isinstance(v, str):
            clean_v = v.strip()
            if clean_v.lower() in {'nan', 'none', 'nat', 'null'}:
                clean_v = ""
        else:
            clean_v = v
        context[sanitize_key(k)] = clean_v
    return context


def _generate_one_pdf(args: tuple):
    """
    Worker function (runs in a separate process):
      1. Fills the .docx template for a single farmer row.
      2. Converts the filled .docx to PDF via LibreOffice or MS Word COM.
      3. Returns (worker_idx, pdf_path | None, error_message | None).

    Each worker gets its own LibreOffice user-profile directory so that
    multiple LO instances can run truly in parallel without conflicts.
    """
    worker_idx, context, template_file, temp_dir_str, lo_profile_base_str, com_lock = args

    temp_dir        = Path(temp_dir_str)
    lo_profile_base = Path(lo_profile_base_str)
    temp_basename   = f"form_{worker_idx:05d}"
    temp_docx_path  = temp_dir / f"{temp_basename}.docx"
    temp_pdf_path   = temp_dir / f"{temp_basename}.pdf"

    # Each worker gets its own isolated LO profile (must be an absolute URI)
    lo_profile_dir = (lo_profile_base / f"lo_{worker_idx:05d}").resolve()
    lo_profile_dir.mkdir(parents=True, exist_ok=True)
    lo_profile_url = lo_profile_dir.as_uri()   # file:///absolute/path/...

    # LibreOffice output log — one per row so conversion errors are inspectable
    lo_log_path = temp_dir / f"{temp_basename}_lo.log"

    try:
        # 1. Render the filled .docx
        doc = DocxTemplate(template_file)
        doc.render(context)
        for section in doc.docx.sections:
            section.page_width  = Mm(210)
            section.page_height = Mm(297)
            
        # Dynamically inject transmittal rows using python-docx to avoid Jinja loop bugs
        if 'farmers' in context and doc.docx.tables:
            table = doc.docx.tables[0]

            # ── Table width ────────────────────────────────────────────────
            # Adjust TABLE_WIDTH_MM to match your page margins.
            # A4 with 2.54 cm margins on each side → usable width ≈ 160 mm.
            TABLE_WIDTH_MM = 160
            table.width = Mm(TABLE_WIDTH_MM)

            # ── Read formatting from the sample row in the template ────────
            # Add a second row (after the header) to Transmittal_Template.docx
            # in Word and format it exactly as you want farmer rows to look
            # (font name, font size, bold, etc.).  The script will copy those
            # settings and apply them to every injected row automatically.
            has_sample_row = len(table.rows) > 1
            sample_row_idx = 1 if has_sample_row else 0
            col_formats = []
            for col_idx in range(len(table.columns)):
                fmt = {'name': None, 'size': None, 'bold': None}
                try:
                    cell = table.cell(sample_row_idx, col_idx)
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        r = cell.paragraphs[0].runs[0]
                        fmt['name'] = r.font.name
                        fmt['size'] = r.font.size
                        fmt['bold'] = r.bold if has_sample_row else False
                except Exception:
                    pass
                col_formats.append(fmt)

            # ── Inject one row per farmer ─────────────────────────────────
            for i, farmer in enumerate(context['farmers']):
                # Reuse the sample row for the first farmer so we don't leave
                # a blank placeholder row in the output.
                if i == 0 and has_sample_row:
                    row_cells = table.rows[1].cells
                else:
                    row_cells = table.add_row().cells

                row_cells[0].text = str(farmer.get('Barangay', ''))
                row_cells[1].text = str(farmer.get('Full_Name', ''))
                row_cells[2].text = str(farmer.get('Birthday', ''))
                row_cells[3].text = str(farmer.get('Gender', ''))
                row_cells[4].text = str(farmer.get('Reference_No', ''))

                # Apply formatting copied from the sample row
                for col_idx, cell in enumerate(row_cells):
                    if col_idx < len(col_formats):
                        fmt = col_formats[col_idx]
                        for p in cell.paragraphs:
                            for r in p.runs:
                                if fmt['name']:
                                    r.font.name = fmt['name']
                                if fmt['size']:
                                    r.font.size = fmt['size']
                                if fmt['bold'] is not None:
                                    r.bold = fmt['bold']
                
        doc.save(temp_docx_path)

        # 2. Convert to PDF (LibreOffice or MS Word COM)
        lo_bin = shutil.which("libreoffice") or shutil.which("soffice")
        if not lo_bin:
            for loc in [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]:
                if os.path.exists(loc):
                    lo_bin = loc
                    break

        if lo_bin:
            with open(lo_log_path, "w") as log_f:
                result = subprocess.run(
                    [
                        lo_bin,
                        f"-env:UserInstallation={lo_profile_url}",
                        "--headless",
                        "--convert-to", "pdf",
                        str(temp_docx_path.resolve()),
                        "--outdir", str(temp_dir.resolve()),
                    ],
                    stdout=log_f,
                    stderr=log_f,
                    check=False,
                )

            if result.returncode == 0 and temp_pdf_path.exists():
                return worker_idx, temp_pdf_path, None

            lo_log = lo_log_path.read_text(errors="replace").strip() if lo_log_path.exists() else ""
            return worker_idx, None, f"LibreOffice rc={result.returncode}:\n{lo_log}"

        if os.name == 'nt':
            if com_lock:
                com_lock.acquire()
            try:
                import win32com.client as win32
                import pythoncom
                pythoncom.CoInitialize()
                word = None
                try:
                    word = win32.DispatchEx("Word.Application")
                    word.Visible = False
                    word.DisplayAlerts = 0
                    doc = word.Documents.Open(str(temp_docx_path.resolve()), ReadOnly=True)
                    doc.SaveAs(str(temp_pdf_path.resolve()), FileFormat=17)
                    doc.Close(0)
                    if temp_pdf_path.exists():
                        return worker_idx, temp_pdf_path, None
                finally:
                    if word:
                        try:
                            word.Quit()
                        except Exception:
                            pass
                    pythoncom.CoUninitialize()
            finally:
                if com_lock:
                    com_lock.release()

        return worker_idx, None, "No PDF engine available (LibreOffice or MS Word required)."

    except Exception as exc:
        return worker_idx, None, str(exc)


def main():
    # ── Directories ────────────────────────────────────────────────────────
    OUTPUT_DIR.mkdir(exist_ok=True)
    TEMP_DIR.mkdir(exist_ok=True)
    LOG_DIR.mkdir(exist_ok=True)

    # ── Logging (console + rotating file) ─────────────────────────────────
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_file  = LOG_DIR / f"run_{timestamp}.log"
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s  %(levelname)-8s  %(message)s",
        datefmt="%H:%M:%S",
        handlers=[
            logging.FileHandler(log_file, encoding="utf-8"),
            logging.StreamHandler(),
        ],
    )
    log = logging.getLogger(__name__)
    log.info(f"Run started — full log at: {log_file}")

    # ── Load Excel data ────────────────────────────────────────────────────
    log.info(f"Loading Excel data from '{EXCEL_FILE}'...")
    df = pd.read_excel(EXCEL_FILE)

    if TEST_ROW_LIMIT is not None:
        log.info(f"TEST MODE — limiting to first {TEST_ROW_LIMIT} rows.")
        df = df.head(TEST_ROW_LIMIT)

    # ── Date formatting ────────────────────────────────────────────────────
    if 'Birthdate' in df.columns:
        df['Birthdate'] = pd.to_datetime(df['Birthdate'], errors='coerce')
    for col in df.select_dtypes(include=['datetime64[ns]']).columns:
        df[col] = df[col].dt.strftime('%m/%d/%Y')

    df = df.fillna("")

    # ── Detect location columns (case-insensitive) ─────────────────────────
    col_map: dict[str, str] = {}
    for col in df.columns:
        key = str(col).strip().lower()
        if key in ("province", "municipality", "barangay"):
            col_map[key] = col

    missing = [c for c in ("province", "municipality", "barangay") if c not in col_map]
    if missing:
        log.error(f"Aborting — required columns not found: {', '.join(missing)}")
        return

    province_col     = col_map["province"]
    municipality_col = col_map["municipality"]
    barangay_col     = col_map["barangay"]

    log.info(f"Total rows to process: {len(df)}")
    log.info("Grouping by Province → Municipality → Barangay...")
    grouped = df.groupby([province_col, municipality_col, barangay_col])

    # ── Build a flat job list (all rows across every barangay) ─────────────
    # Processing all rows in parallel rather than barangay-by-barangay
    # dramatically reduces wall-clock time.
    jobs:      list[tuple] = []
    job_meta:  list[tuple] = []   # (province, municipality, barangay, is_transmittal) per job
    lo_profile_base = str(TEMP_DIR.resolve() / "lo_profiles")

    import multiprocessing
    manager = multiprocessing.Manager()
    com_lock = manager.Lock()

    for (province_name, municipality_name, barangay_name), group_df in grouped:
        if not province_name or not municipality_name or not barangay_name:
            continue  # skip fully-empty location entries

        # Sort farmers alphabetically by Full Name
        if 'Full Name' in group_df.columns:
            group_df = group_df.sort_values(by='Full Name')

        # 1. Prepare Transmittal Data
        farmers_list = []
        for _, row in group_df.iterrows():
            ctx = _build_context(row.to_dict())
            farmers_list.append({
                'Barangay': barangay_name,
                'Full_Name': ctx.get('Full_Name', ''),
                'Birthday': ctx.get('Birthdate', ''),
                'Gender': ctx.get('Gender', ''),
                'Reference_No': ctx.get('Reference_No', '')
            })

        # 2. Submit Transmittal Job
        transmittal_idx = len(jobs)
        jobs.append((
            transmittal_idx,
            {'barangay': f"{barangay_name}, {municipality_name}", 'farmers': farmers_list},
            TRANSMITTAL_TEMPLATE,
            str(TEMP_DIR.resolve()),
            lo_profile_base,
            com_lock,
        ))
        job_meta.append((province_name, municipality_name, barangay_name, True))

        # 3. Submit Individual Form Jobs
        for enum_idx, (_, row) in enumerate(group_df.iterrows()):
            worker_idx = len(jobs)
            jobs.append((
                worker_idx,
                _build_context(row.to_dict()),
                TEMPLATE_FILE,
                str(TEMP_DIR.resolve()),
                lo_profile_base,
                com_lock,
            ))
            job_meta.append((province_name, municipality_name, barangay_name, False))

    log.info(f"Submitting {len(jobs)} rows to {MAX_WORKERS} parallel workers...")

    # ── Parallel PDF generation ────────────────────────────────────────────
    pdf_results: dict[int, Path | None] = {}
    error_count = 0

    with concurrent.futures.ProcessPoolExecutor(max_workers=MAX_WORKERS) as pool:
        future_to_idx = {pool.submit(_generate_one_pdf, job): job[0] for job in jobs}
        done = 0
        for future in concurrent.futures.as_completed(future_to_idx):
            worker_idx, pdf_path, error_msg = future.result()
            done += 1
            if error_msg:
                log.warning(f"  [{done}/{len(jobs)}] Row {worker_idx} FAILED: {error_msg}")
                error_count += 1
            else:
                log.info(f"  [{done}/{len(jobs)}] Row {worker_idx} OK → {pdf_path.name}")
            pdf_results[worker_idx] = pdf_path

    log.info(
        f"Generation complete — "
        f"{len(jobs) - error_count}/{len(jobs)} PDFs OK, {error_count} failed."
    )

    # ── Merge PDFs grouped by Barangay ─────────────────────────────────────
    barangay_transmittals: dict[tuple, Path] = {}
    barangay_forms: dict[tuple, list[Path]] = defaultdict(list)
    
    for worker_idx, meta in enumerate(job_meta):
        prov, muni, brgy, is_trans = meta
        group_key = (prov, muni, brgy)
        pdf_path = pdf_results.get(worker_idx)
        if pdf_path:
            if is_trans:
                barangay_transmittals[group_key] = pdf_path
            else:
                barangay_forms[group_key].append(pdf_path)

    total_merged = 0
    for group_key, form_list in barangay_forms.items():
        province_name, municipality_name, barangay_name = group_key
        safe_prov  = safe_filename(province_name)
        safe_muni  = safe_filename(municipality_name)
        safe_brgy  = safe_filename(barangay_name)

        folder_name = f"{safe_prov}, {safe_muni}"
        target_dir  = OUTPUT_DIR / folder_name
        target_dir.mkdir(parents=True, exist_ok=True)

        merged_path = target_dir / f"Barangay_{safe_brgy}_{len(form_list)}_Forms.pdf"
        log.info(f"  Merging transmittal + {len(form_list):3d} forms → {merged_path}")

        merger = PdfWriter()
        
        # Append transmittal first
        transmittal_pdf = barangay_transmittals.get(group_key)
        if transmittal_pdf:
            merger.append(str(transmittal_pdf))
            
        # Append forms
        for pdf_path in form_list:
            merger.append(str(pdf_path))
            
        with open(merged_path, "wb") as f:
            merger.write(f)

        total_merged += len(form_list)

    # ── Cleanup ────────────────────────────────────────────────────────────
    if error_count == 0:
        log.info("No errors — removing temporary files...")
        shutil.rmtree(TEMP_DIR)
    else:
        log.warning(
            f"{error_count} row(s) failed. "
            f"Temp files preserved for inspection at: {TEMP_DIR.resolve()}"
        )

    log.info(f"All done! {total_merged} PDFs merged into '{OUTPUT_DIR}/'.")


if __name__ == "__main__":
    main()
