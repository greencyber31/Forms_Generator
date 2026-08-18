import pandas as pd
from docxtpl import DocxTemplate
from docx.shared import Mm
import os
import re
import subprocess
import concurrent.futures
import logging
import shutil
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from pypdf import PdfWriter
import traceback

_UNSAFE_CHARS = str.maketrans({
    '/': '_', '\\': '_', ':': '_', '*': '_',
    '?': '_', '"': '_', '<': '_', '>': '_', '|': '_',
})

def sanitize_key(k: str) -> str:
    """Convert an Excel column header into a valid Jinja2 variable name."""
    s = str(k).strip()
    s = s.replace(' ', '_').replace('.', '').replace('-', '_').replace('/', '_')
    return ''.join(c for c in s if c.isalnum() or c == '_')

def safe_filename(name: str) -> str:
    """Replace all OS-reserved / path-unsafe characters in a filename segment."""
    return str(name).translate(_UNSAFE_CHARS).strip()

def _build_context(row_dict: dict, mapping: dict | None = None) -> dict:
    """Convert a raw row dictionary into a clean Jinja2 rendering context using mapped tags."""
    context = {}
    for k, v in row_dict.items():
        if pd.isna(v):
            clean_v = ""
        elif isinstance(v, str):
            clean_v = v.strip()
            if clean_v.lower() in {'nan', 'none', 'nat', 'null'}:
                clean_v = ""
        else:
            clean_v = v
        context[sanitize_key(k)] = clean_v

    if mapping:
        for tag, excel_col in mapping.items():
            if not excel_col:
                continue
            if isinstance(excel_col, str) and excel_col.startswith("STATIC:"):
                context[tag] = excel_col[7:]
            elif excel_col in row_dict:
                val = row_dict[excel_col]
                if pd.isna(val):
                    clean_v = ""
                elif isinstance(val, str):
                    clean_v = val.strip()
                    if clean_v.lower() in {'nan', 'none', 'nat', 'null'}:
                        clean_v = ""
                else:
                    clean_v = val
                context[tag] = clean_v

    return context

def _ensure_table_borders(table):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tblPr = table._tbl.tblPr
        tblBorders = tblPr.first_child_found_in("w:tblBorders")
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)

        borders = {
            'top': {'val': 'single', 'sz': '4', 'space': '0', 'color': '000000'},
            'left': {'val': 'single', 'sz': '4', 'space': '0', 'color': '000000'},
            'bottom': {'val': 'single', 'sz': '4', 'space': '0', 'color': '000000'},
            'right': {'val': 'single', 'sz': '4', 'space': '0', 'color': '000000'},
            'insideH': {'val': 'single', 'sz': '4', 'space': '0', 'color': '000000'},
            'insideV': {'val': 'single', 'sz': '4', 'space': '0', 'color': '000000'}
        }
        for border_name, border_props in borders.items():
            border_el = tblBorders.find(qn(f'w:{border_name}'))
            if border_el is None:
                border_el = OxmlElement(f'w:{border_name}')
                tblBorders.append(border_el)
            for k, v in border_props.items():
                border_el.set(qn(f'w:{k}'), str(v))
    except Exception:
        pass

def _get_farmer_val(farmer: dict, tag_key: str, mapping: dict | None = None) -> str:
    if not tag_key or not farmer:
        return ""

    mapped_col = mapping.get(tag_key) if mapping else None
    if mapped_col:
        if isinstance(mapped_col, str) and mapped_col.startswith("STATIC:"):
            return mapped_col[7:]
        val = farmer.get(mapped_col, farmer.get(sanitize_key(mapped_col)))
        if val is not None and not pd.isna(val) and str(val).lower() not in {'nan', 'none', 'null'}:
            return str(val)

    if tag_key in farmer and farmer[tag_key] is not None and not pd.isna(farmer[tag_key]):
        return str(farmer[tag_key])

    tag_clean = sanitize_key(tag_key).lower()
    farmer_lower = {sanitize_key(k).lower(): v for k, v in farmer.items() if v is not None and not pd.isna(v)}
    if tag_clean in farmer_lower:
        return str(farmer_lower[tag_clean])

    if 'name' in tag_clean:
        if 'last' in tag_clean and 'last_name' in farmer_lower:
            return str(farmer_lower['last_name'])
        if 'first' in tag_clean and 'first_name' in farmer_lower:
            return str(farmer_lower['first_name'])
        if 'middle' in tag_clean and 'middle_name' in farmer_lower:
            return str(farmer_lower['middle_name'])
        if 'full_name' in farmer_lower:
            return str(farmer_lower['full_name'])

    if 'ref' in tag_clean or 'id' in tag_clean or 'no' in tag_clean:
        for r_key in ['rsbsa_no', 'reference_no', 'reference_no.', 'id', 'farmer_id', 'ref_no']:
            if r_key in farmer_lower:
                return str(farmer_lower[r_key])

    if 'birth' in tag_clean or 'bday' in tag_clean:
        for b_key in ['birthdate', 'birth_date', 'birthday', 'bday']:
            if b_key in farmer_lower:
                return str(farmer_lower[b_key])

    if 'gender' in tag_clean or 'sex' in tag_clean:
        for g_key in ['gender', 'sex']:
            if g_key in farmer_lower:
                return str(farmer_lower[g_key])

    if 'brgy' in tag_clean or 'barangay' in tag_clean:
        for br_key in ['barangay', 'brgy', 'home_barangay', 'farm_barangay']:
            if br_key in farmer_lower:
                return str(farmer_lower[br_key])

    return ""

def _calculate_smart_column_widths(active_cols: list[dict], farmers_list: list[dict], total_width_mm: float = 160.0) -> list[float]:
    """
    Calculates content-aware column widths based on header length and max data row text length.
    Ensures short columns (Gender, Suffix, ID) take less space and long columns (Names, Address, Contact) get more space.
    Strictly normalizes all column widths so their sum is EXACTLY equal to total_width_mm (160mm),
    preventing any table overflow beyond the A4 page margins.
    """
    if not active_cols:
        return []

    scores = []
    for col_info in active_cols:
        h_name = str(col_info.get("header", ""))
        max_char_len = max(len(h_name), 4)

        for farmer in farmers_list[:50]:
            val = _get_farmer_val(farmer, h_name, None)
            if val:
                max_char_len = max(max_char_len, min(len(str(val)), 30))

        scores.append(max_char_len ** 0.85)

    total_score = sum(scores)
    if total_score <= 0:
        total_score = 1.0

    raw_widths = [(score / total_score) * total_width_mm for score in scores]
    min_col_mm = max(10.0, total_width_mm / (len(active_cols) * 2.5))
    floored_widths = [max(w, min_col_mm) for w in raw_widths]

    floored_sum = sum(floored_widths)
    final_widths = [(w / floored_sum) * total_width_mm for w in floored_widths]

    return final_widths

def _populate_transmittal_table(table, farmers_list: list[dict], transmittal_mapping: dict | None = None, transmittal_columns: list[dict] | None = None):
    if not farmers_list:
        return

    header_row_idx = 0
    sample_row_idx = 1 if len(table.rows) > 1 else 0

    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        row0_texts = set([c.text.strip() for c in table.rows[0].cells])
        if len(row0_texts) == 1 and len(table.rows) > 2:
            header_row_idx = 1
            sample_row_idx = 2
            header_row_indices = [0, 1]
        else:
            header_row_indices = [0]

        for r_idx in header_row_indices:
            trPr = table.rows[r_idx]._tr.get_or_add_trPr()
            if trPr.find(qn('w:tblHeader')) is None:
                trPr.append(OxmlElement('w:tblHeader'))
    except Exception:
        pass

    # Check for active transmittal columns
    active_cols = []
    if transmittal_columns:
        active_cols = [c for c in transmittal_columns if c.get('enabled', True)]
        def safe_order(c):
            val = c.get('order')
            try:
                return int(val)
            except Exception:
                return 999
        active_cols.sort(key=safe_order)

    if active_cols:
        target_count = len(active_cols)
        col_widths_mm = _calculate_smart_column_widths(active_cols, farmers_list, total_width_mm=160.0)

        while len(table.columns) < target_count:
            table.add_column(Mm(col_widths_mm[min(len(table.columns), target_count - 1)]))

        header_cells = table.rows[header_row_idx].cells
        for col_i, col_info in enumerate(active_cols):
            if col_i < len(header_cells):
                header_cells[col_i].text = str(col_info["header"])

        has_sample_row = len(table.rows) > sample_row_idx
        col_formats = []
        for col_idx in range(target_count):
            fmt = {'name': None, 'size': None, 'bold': None}
            try:
                cell = table.cell(sample_row_idx, col_idx)
                if cell.paragraphs and cell.paragraphs[0].runs:
                    r = cell.paragraphs[0].runs[0]
                    fmt['name'] = r.font.name
                    fmt['size'] = r.font.size
                    fmt['bold'] = r.bold if has_sample_row else False
            except Exception:
                pass
            col_formats.append(fmt)

        for i, farmer in enumerate(farmers_list):
            if i == 0 and has_sample_row:
                row_cells = table.rows[sample_row_idx].cells
            else:
                row_cells = table.add_row().cells

            for col_i, col_info in enumerate(active_cols):
                if col_i >= len(row_cells):
                    break
                header_name = col_info["header"]
                val = _get_farmer_val(farmer, header_name, transmittal_mapping)
                row_cells[col_i].text = str(val) if val is not None else ""

            for col_idx, cell in enumerate(row_cells):
                if col_idx < len(col_formats):
                    fmt = col_formats[col_idx]
                    for p in cell.paragraphs:
                        for r in p.runs:
                            if fmt['name']:
                                r.font.name = fmt['name']
                            if fmt['size']:
                                r.font.size = fmt['size']
                            if fmt['bold'] is not None:
                                r.bold = fmt['bold']

        # Enforce exact smart width on every row & cell to maintain strict A4 margins (160mm total)
        for row in table.rows:
            for col_i, w_mm in enumerate(col_widths_mm):
                if col_i < len(row.cells):
                    row.cells[col_i].width = Mm(w_mm)

        _ensure_table_borders(table)
        return

    col_tags = {}
    header_cells = table.rows[header_row_idx].cells
    for col_i, cell in enumerate(header_cells):
        raw_text = cell.text.strip()
        matches = re.findall(r'\{\{\s*([^\}\s]+)\s*\}\}', raw_text)
        if matches:
            tag_name = matches[0]
            col_tags[col_i] = tag_name
            mapped_val = transmittal_mapping.get(tag_name) if transmittal_mapping else None
            if mapped_val:
                display_label = mapped_val[7:] if mapped_val.startswith("STATIC:") else mapped_val
            else:
                display_label = tag_name.replace('_', ' ').title()
            cell.text = display_label
        else:
            col_tags[col_i] = raw_text

    has_sample_row = len(table.rows) > sample_row_idx
    col_formats = []
    for col_idx in range(len(table.columns)):
        fmt = {'name': None, 'size': None, 'bold': None}
        try:
            cell = table.cell(sample_row_idx, col_idx)
            if cell.paragraphs and cell.paragraphs[0].runs:
                r = cell.paragraphs[0].runs[0]
                fmt['name'] = r.font.name
                fmt['size'] = r.font.size
                fmt['bold'] = r.bold if has_sample_row else False
        except Exception:
            pass
        col_formats.append(fmt)

    for i, farmer in enumerate(farmers_list):
        if i == 0 and has_sample_row:
            row_cells = table.rows[sample_row_idx].cells
        else:
            row_cells = table.add_row().cells

        for col_i in range(len(table.columns)):
            if col_i >= len(row_cells):
                break
            tag_key = col_tags.get(col_i, "")
            val = _get_farmer_val(farmer, tag_key, transmittal_mapping)
            row_cells[col_i].text = str(val) if val is not None else ""

        for col_idx, cell in enumerate(row_cells):
            if col_idx < len(col_formats):
                fmt = col_formats[col_idx]
                for p in cell.paragraphs:
                    for r in p.runs:
                        if fmt['name']:
                            r.font.name = fmt['name']
                        if fmt['size']:
                            r.font.size = fmt['size']
                        if fmt['bold'] is not None:
                            r.bold = fmt['bold']

    _ensure_table_borders(table)

def _generate_one_pdf(args: tuple):
    """
    Worker function (runs in a separate process):
      1. Fills the .docx template for a single farmer row or transmittal.
      2. Converts the filled .docx to PDF via LibreOffice (parallel) or MS Word COM (process locked).
      3. Returns (worker_idx, pdf_path | None, error_message | None).
    """
    worker_idx, context, template_file, temp_dir_str, lo_profile_base_str, com_lock = args

    temp_dir        = Path(temp_dir_str)
    lo_profile_base = Path(lo_profile_base_str)
    temp_basename   = f"form_{worker_idx:05d}"
    temp_docx_path  = temp_dir / f"{temp_basename}.docx"
    temp_pdf_path   = temp_dir / f"{temp_basename}.pdf"

    # Profile isolation for parallel LibreOffice instances
    lo_profile_dir = (lo_profile_base / f"lo_{worker_idx:05d}").resolve()
    lo_profile_dir.mkdir(parents=True, exist_ok=True)
    lo_profile_url = lo_profile_dir.as_uri()

    lo_log_path = temp_dir / f"{temp_basename}_lo.log"

    try:
        doc = DocxTemplate(template_file)
        doc.init_docx()

        # Transmittal list table injection if farmers array is present (BEFORE doc.render)
        if 'farmers' in context and doc.docx.tables:
            _populate_transmittal_table(doc.docx.tables[0], context['farmers'], context.get('_transmittal_mapping'), context.get('_transmittal_columns'))

        doc.render(context)
        
        # Enforce A4 page bounds on sections
        for section in doc.docx.sections:
            section.page_width  = Mm(210)
            section.page_height = Mm(297)
            
        # Transmittal list table injection if farmers array is present
        if 'farmers' in context and doc.docx.tables:
            _populate_transmittal_table(doc.docx.tables[0], context['farmers'], context.get('_transmittal_mapping'), context.get('_transmittal_columns'))

        doc.save(temp_docx_path)

        # Convert .docx to PDF (LibreOffice or MS Word COM)
        lo_bin = shutil.which("libreoffice") or shutil.which("soffice")
        if not lo_bin:
            for loc in [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]:
                if os.path.exists(loc):
                    lo_bin = loc
                    break

        if lo_bin:
            with open(lo_log_path, "w") as log_f:
                result = subprocess.run(
                    [
                        lo_bin,
                        f"-env:UserInstallation={lo_profile_url}",
                        "--headless",
                        "--convert-to", "pdf",
                        str(temp_docx_path.resolve()),
                        "--outdir", str(temp_dir.resolve()),
                    ],
                    stdout=log_f,
                    stderr=log_f,
                    check=False,
                )
            if result.returncode == 0 and temp_pdf_path.exists():
                return worker_idx, temp_pdf_path, None
            lo_log = lo_log_path.read_text(errors="replace").strip() if lo_log_path.exists() else ""
            return worker_idx, None, f"LibreOffice rc={result.returncode}:\n{lo_log}"

        if os.name == 'nt':
            if com_lock:
                com_lock.acquire()
            try:
                import win32com.client as win32
                import pythoncom
                pythoncom.CoInitialize()
                word = None
                try:
                    word = win32.DispatchEx("Word.Application")
                    word.Visible = False
                    word.DisplayAlerts = 0
                    doc = word.Documents.Open(str(temp_docx_path.resolve()), ReadOnly=True)
                    doc.SaveAs(str(temp_pdf_path.resolve()), FileFormat=17)
                    doc.Close(0)
                    if temp_pdf_path.exists():
                        return worker_idx, temp_pdf_path, None
                finally:
                    if word:
                        try:
                            word.Quit()
                        except Exception:
                            pass
                    pythoncom.CoUninitialize()
            finally:
                if com_lock:
                    com_lock.release()

        return worker_idx, None, "No PDF engine available (LibreOffice or MS Word required)."

    except Exception as exc:
        return worker_idx, None, str(exc)


def run_batch_generation(
    excel_file: str,
    template_file: str,
    transmittal_template: str,
    output_dir: str = "output",
    temp_dir: str = "temp_docs",
    primary_group_col: str | None = None,
    secondary_group_col: str | None = None,
    bundle_group_col: str | None = None,
    test_limit: int | None = None,
    max_workers: int = 4,
    template_mapping: dict | None = None,
    transmittal_mapping: dict | None = None,
    transmittal_columns: list[dict] | None = None,
    progress_callback=None
):
    """
    Executes full batch generation and returns execution summary.
    Emits real-time SSE progress events if progress_callback is provided.
    """
    out_path = Path(output_dir)
    tmp_path = Path(temp_dir)
    out_path.mkdir(parents=True, exist_ok=True)
    tmp_path.mkdir(parents=True, exist_ok=True)

    def log(msg, status="info"):
        if progress_callback:
            progress_callback({
                "type": "log",
                "status": status,
                "message": msg,
                "timestamp": datetime.now().strftime("%H:%M:%S")
            })

    log(f"Starting batch generation from '{Path(excel_file).name}'")
    df = pd.read_excel(excel_file)

    if test_limit and test_limit > 0:
        log(f"Test Mode: Limiting to first {test_limit} rows.")
        df = df.head(test_limit)

    # Date column formatting
    for col in df.columns:
        if 'date' in str(col).lower() or 'birth' in str(col).lower():
            df[col] = pd.to_datetime(df[col], errors='coerce')
    for col in df.select_dtypes(include=['datetime64[ns]']).columns:
        df[col] = df[col].dt.strftime('%m/%d/%Y')

    df = df.fillna("")

    # Resolve case-insensitive column matching
    col_dict = {str(c).strip().lower(): str(c) for c in df.columns}

    def resolve_col(user_choice, default_names):
        if user_choice and user_choice.strip().lower() in col_dict:
            return col_dict[user_choice.strip().lower()]
        for d in default_names:
            if d.lower() in col_dict:
                return col_dict[d.lower()]
        return None

    p_col = resolve_col(primary_group_col, ["province", "region", "district"])
    s_col = resolve_col(secondary_group_col, ["municipality", "city", "town"])
    b_col = resolve_col(bundle_group_col, ["barangay", "association", "group", "cooperative"])

    group_keys = [c for c in [p_col, s_col, b_col] if c is not None]

    if not group_keys:
        log("No valid grouping columns found. Treating dataset as single bundle.", "warning")
        grouped = [("All_Records", df)]
    else:
        log(f"Grouping dataset by: {' > '.join(group_keys)}")
        grouped = df.groupby(group_keys)

    import multiprocessing
    manager = multiprocessing.Manager()
    com_lock = manager.Lock()

    jobs: list[tuple] = []
    job_meta: list[tuple] = []
    lo_profile_base = str((tmp_path / "lo_profiles").resolve())

    for group_key_val, group_df in grouped:
        if not isinstance(group_key_val, tuple):
            group_key_tuple = (group_key_val,)
        else:
            group_key_tuple = group_key_val

        # Extract names safely
        prov_name = str(group_key_tuple[0]) if len(group_key_tuple) > 0 else ""
        muni_name = str(group_key_tuple[1]) if len(group_key_tuple) > 1 else ""
        brgy_name = str(group_key_tuple[2]) if len(group_key_tuple) > 2 else prov_name

        if not any(group_key_tuple):
            continue

        # Sort farmers by Full Name if column exists
        if 'Full Name' in group_df.columns:
            group_df = group_df.sort_values(by='Full Name')

        # 1. Transmittal Data
        farmers_list = []
        for _, row in group_df.iterrows():
            r_d = row.to_dict()
            ctx = _build_context(r_d, transmittal_mapping)
            ctx['Barangay'] = brgy_name
            if 'Full_Name' not in ctx:
                ctx['Full_Name'] = ctx.get('Fullname', ctx.get('Name', ''))
            if 'Birthday' not in ctx:
                ctx['Birthday'] = ctx.get('Birthdate', ctx.get('Birth_Date', ''))
            if 'Reference_No' not in ctx:
                ctx['Reference_No'] = ctx.get('Reference_No.', ctx.get('ID', ''))
            # Store all original raw row fields as fallbacks
            for r_k, r_v in r_d.items():
                if r_k not in ctx:
                    ctx[r_k] = r_v
            farmers_list.append(ctx)

        # 2. Submit Transmittal Job
        transmittal_idx = len(jobs)
        transmittal_ctx = {
            'barangay': brgy_name,
            'municipality': muni_name,
            'province': prov_name,
            'farmers': farmers_list,
            '_transmittal_mapping': transmittal_mapping,
            '_transmittal_columns': transmittal_columns
        }
        if transmittal_mapping:
            for tag, val in transmittal_mapping.items():
                if val.startswith('STATIC:'):
                    transmittal_ctx[tag] = val[7:]
        jobs.append((
            transmittal_idx,
            transmittal_ctx,
            transmittal_template,
            str(tmp_path.resolve()),
            lo_profile_base,
            com_lock,
        ))
        job_meta.append((prov_name, muni_name, brgy_name, True))

        # 3. Submit Individual Form Jobs
        for _, row in group_df.iterrows():
            worker_idx = len(jobs)
            jobs.append((
                worker_idx,
                _build_context(row.to_dict(), template_mapping),
                template_file,
                str(tmp_path.resolve()),
                lo_profile_base,
                com_lock,
            ))
            job_meta.append((prov_name, muni_name, brgy_name, False))

    log(f"Submitting {len(jobs)} rendering jobs across {max_workers} parallel workers...")

    pdf_results: dict[int, Path | None] = {}
    error_count = 0
    total_jobs = len(jobs)

    with concurrent.futures.ProcessPoolExecutor(max_workers=max_workers) as pool:
        future_to_idx = {pool.submit(_generate_one_pdf, job): job[0] for job in jobs}
        done = 0
        for future in concurrent.futures.as_completed(future_to_idx):
            worker_idx, pdf_path, error_msg = future.result()
            done += 1
            if error_msg:
                log(f"Job [{done}/{total_jobs}] FAILED: {error_msg}", "warning")
                error_count += 1
            else:
                pdf_results[worker_idx] = pdf_path

            if progress_callback:
                progress_callback({
                    "type": "progress",
                    "current": done,
                    "total": total_jobs,
                    "percent": int((done / total_jobs) * 100),
                    "failed": error_count
                })

    log(f"Generation complete — {total_jobs - error_count}/{total_jobs} PDFs rendered successfully.")

    # PDF Merging
    bundle_transmittals: dict[tuple, Path] = {}
    bundle_forms: dict[tuple, list[Path]] = defaultdict(list)

    for worker_idx, meta in enumerate(job_meta):
        prov, muni, brgy, is_trans = meta
        group_key = (prov, muni, brgy)
        pdf_path = pdf_results.get(worker_idx)
        if pdf_path:
            if is_trans:
                bundle_transmittals[group_key] = pdf_path
            else:
                bundle_forms[group_key].append(pdf_path)

    total_merged = 0
    for group_key, form_list in bundle_forms.items():
        prov, muni, brgy = group_key
        safe_prov = safe_filename(prov)
        safe_muni = safe_filename(muni)
        safe_brgy = safe_filename(brgy)

        # Build output directory path
        if safe_prov and safe_muni:
            target_dir = out_path / f"{safe_prov}, {safe_muni}"
        elif safe_prov:
            target_dir = out_path / safe_prov
        else:
            target_dir = out_path

        target_dir.mkdir(parents=True, exist_ok=True)

        merged_path = target_dir / f"Bundle_{safe_brgy}_{len(form_list)}_Forms.pdf"
        log(f"Merging transmittal + {len(form_list)} forms into -> {merged_path.name}")

        merger = PdfWriter()
        trans_pdf = bundle_transmittals.get(group_key)
        if trans_pdf:
            merger.append(str(trans_pdf))

        for p_file in form_list:
            merger.append(str(p_file))

        with open(merged_path, "wb") as f:
            merger.write(f)

        total_merged += len(form_list)

    # Cleanup temp directory if no errors
    if error_count == 0:
        log("No errors recorded — cleaning up temporary files.")
        shutil.rmtree(tmp_path, ignore_errors=True)

    log(f"Batch execution finished! {total_merged} application forms merged into '{out_path}/'.", "success")

    return {
        "total_rendered": total_jobs - error_count,
        "total_failed": error_count,
        "total_merged": total_merged,
        "output_directory": str(out_path.resolve())
    }
