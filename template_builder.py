import re
import os
import subprocess
import shutil
import uuid
from pathlib import Path
import pandas as pd
import docx
from docxtpl import DocxTemplate

def sanitize_key(k: str) -> str:
    """Convert an Excel column header into a valid Jinja2 variable name."""
    s = str(k).strip()
    s = s.replace(' ', '_').replace('.', '').replace('-', '_').replace('/', '_')
    return ''.join(c for c in s if c.isalnum() or c == '_')

def extract_excel_headers(excel_path: str) -> list[dict]:
    """Reads Excel file and returns column metadata with sanitized Jinja keys."""
    try:
        df = pd.read_excel(excel_path, nrows=2)
        headers = []
        for col in df.columns:
            orig = str(col).strip()
            sanitized = sanitize_key(orig)
            headers.append({
                "original": orig,
                "sanitized": sanitized,
                "tag": f"{{{{ {sanitized} }}}}"
            })
        return headers
    except Exception as e:
        raise ValueError(f"Failed to read Excel headers: {str(e)}")

def extract_docx_tags(docx_path: str) -> list[str]:
    """Inspects a Word document and extracts all {{ variable }} tags in top-to-bottom document order."""
    try:
        doc = docx.Document(docx_path)
        found_tags = []
        seen = set()

        def extract_text_tags(text: str):
            matches = re.findall(r'\{\{\s*([^\}\s]+)\s*\}\}', text)
            for tag in matches:
                clean_tag = sanitize_key(tag)
                if clean_tag and clean_tag not in seen:
                    seen.add(clean_tag)
                    found_tags.append(clean_tag)

        for element in doc.element.body:
            if element.tag.endswith('p'):
                p = docx.text.paragraph.Paragraph(element, doc)
                extract_text_tags(p.text)
            elif element.tag.endswith('tbl'):
                table = docx.table.Table(element, doc)
                for row in table.rows:
                    for cell in row.cells:
                        for cell_p in cell.paragraphs:
                            extract_text_tags(cell_p.text)

        return found_tags
    except Exception as e:
        raise ValueError(f"Failed to inspect docx tags: {str(e)}")

def auto_match_tags(excel_headers: list[str], docx_tags: list[str]) -> dict[str, str]:
    """
    Intelligently connects Word document Jinja tags to Excel column headers.
    Returns a dictionary mapping { docx_tag: excel_header_name }.
    """
    mapping = {}
    
    # Custom explicit mappings for standard insurance fields
    aliases = {
        'birthdate': ['birth date', 'bday', 'birthdate', 'dob'],
        'contact_number': ['contact no', 'contact number', 'phone', 'mobile'],
        'home_barangay': ['barangay', 'brgy'],
        'home_municipality': ['municipality', 'city', 'town'],
        'home_province': ['province'],
        'home_sitiopurok': ['sitio', 'purok', 'street'],
        'farm_barangay': ['barangay', 'brgy'],
        'farm_municipality': ['municipality'],
        'farm_province': ['province'],
        'farm_sitiopurok': ['sitio', 'purok'],
        'cfitf_no': ['reference no', 'reference no.', 'cfitf no'],
        'rsbsa_no': ['farmer id', 'rsbsa no'],
        'area': ['coconutarea', 'area'],
        'pb_bday': ['birth date'],
        'pb_rel': ['civil status']
    }

    for tag in docx_tags:
        t_clean = tag.lower().strip()
        matched = False
        
        # 1. Check exact match
        for h in excel_headers:
            h_clean = h.lower().strip()
            if t_clean == h_clean or t_clean == h_clean.replace(' ', '_'):
                mapping[tag] = h
                matched = True
                break
        if matched:
            continue

        # 2. Check aliases
        if t_clean in aliases:
            for alias in aliases[t_clean]:
                for h in excel_headers:
                    if alias in h.lower():
                        mapping[tag] = h
                        matched = True
                        break
                if matched:
                    break
        if matched:
            continue

        # 3. Check fuzzy containment
        t_norm = t_clean.replace('_', '').replace(' ', '')
        for h in excel_headers:
            h_norm = h.lower().replace('_', '').replace(' ', '').replace('.', '')
            if t_norm == h_norm or t_norm in h_norm or h_norm in t_norm:
                mapping[tag] = h
                matched = True
                break

    return mapping

def render_docx_to_pdf_preview(docx_path: str, temp_dir: str, output_name: str) -> str:
    """Converts docx to high-fidelity PDF via LibreOffice or MS Word for 1:1 exact visual preview."""
    try:
        out_dir = Path(temp_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        pdf_path = out_dir / f"{output_name}.pdf"
        docx_abs = os.path.abspath(docx_path)
        pdf_abs = os.path.abspath(str(pdf_path))

        lo_bin = shutil.which("libreoffice") or shutil.which("soffice")
        if not lo_bin:
            for loc in [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]:
                if os.path.exists(loc):
                    lo_bin = loc
                    break

        if lo_bin:
            profile_dir = (out_dir / f"profile_{output_name}").resolve()
            profile_dir.mkdir(parents=True, exist_ok=True)
            profile_url = profile_dir.as_uri()

            subprocess.run(
                [
                    lo_bin,
                    f"-env:UserInstallation={profile_url}",
                    "--headless",
                    "--convert-to", "pdf",
                    docx_abs,
                    "--outdir", str(out_dir),
                ],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=True
            )

            generated_pdf = out_dir / f"{Path(docx_path).stem}.pdf"
            if generated_pdf.exists() and generated_pdf.resolve() != pdf_path.resolve():
                if pdf_path.exists():
                    try:
                        os.remove(pdf_path)
                    except Exception:
                        pass
                os.rename(generated_pdf, pdf_path)

            return str(pdf_path)

        if os.name == 'nt':
            import win32com.client as win32
            import pythoncom
            pythoncom.CoInitialize()
            word = None
            try:
                word = win32.DispatchEx("Word.Application")
                word.Visible = False
                word.DisplayAlerts = 0
                doc = word.Documents.Open(docx_abs, ReadOnly=True)
                doc.SaveAs(pdf_abs, FileFormat=17)
                doc.Close(0)
                if os.path.exists(pdf_abs):
                    return pdf_abs
            finally:
                if word:
                    try:
                        word.Quit()
                    except Exception:
                        pass
                pythoncom.CoUninitialize()

        raise RuntimeError("No PDF engine available (LibreOffice or MS Word required).")
    except Exception as e:
        raise ValueError(f"Failed to convert docx to preview PDF: {str(e)}")

def render_sample_pdf_preview(
    excel_path: str,
    docx_path: str,
    mapping: dict[str, str],
    temp_dir: str,
    output_name: str,
    transmittal_columns: list[dict] | None = None
) -> str:
    """
    Renders sample Row 1 data from Excel into the Word template using docxtpl,
    then converts to PDF for live 1:1 previewing in the browser.
    """
    try:
        docx_tags = extract_docx_tags(docx_path)
        context = {tag: "" for tag in docx_tags}

        row_dict = {}
        if excel_path and os.path.exists(excel_path):
            try:
                df = pd.read_excel(excel_path, nrows=1)
                if not df.empty:
                    row_dict = df.iloc[0].to_dict()
            except Exception:
                pass

        if row_dict:
            from generator_engine import _build_context
            b_ctx = _build_context(row_dict, mapping)
            context.update(b_ctx)
        else:
            for tag in docx_tags:
                mapped_val = mapping.get(tag)
                if mapped_val and isinstance(mapped_val, str) and mapped_val.startswith("STATIC:"):
                    context[tag] = mapped_val[7:]

        tpl = DocxTemplate(docx_path)
        tpl.init_docx()

        # Only populate transmittal summary table if this is a transmittal template preview (BEFORE render)
        if tpl.docx.tables and 'transmittal' in output_name.lower():
            farmers_list = []
            if excel_path and os.path.exists(excel_path):
                try:
                    df = pd.read_excel(excel_path, nrows=5)
                    from generator_engine import _build_context
                    for _, r in df.iterrows():
                        r_d = r.to_dict()
                        f_ctx = _build_context(r_d, mapping)
                        for r_k, r_v in r_d.items():
                            if r_k not in f_ctx:
                                f_ctx[r_k] = r_v
                        farmers_list.append(f_ctx)
                except Exception:
                    pass

            if not farmers_list:
                farmers_list = [
                    {'Barangay': 'Sample Brgy', 'Full_Name': 'DELA CRUZ, JUAN', 'Birthday': '01/01/1990', 'Gender': 'Male', 'Reference_No': 'REF-001'},
                    {'Barangay': 'Sample Brgy', 'Full_Name': 'SANTOS, MARIA', 'Birthday': '02/02/1992', 'Gender': 'Female', 'Reference_No': 'REF-002'}
                ]

            try:
                from generator_engine import _populate_transmittal_table
                _populate_transmittal_table(tpl.docx.tables[0], farmers_list, mapping, transmittal_columns)
            except Exception:
                pass

        tpl.render(context)
        
        out_dir = Path(temp_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        unique_id = uuid.uuid4().hex[:8]
        sample_docx_name = f"{output_name}_{unique_id}_sample.docx"
        filled_docx = out_dir / sample_docx_name
        tpl.save(filled_docx)

        pdf_path = render_docx_to_pdf_preview(str(filled_docx), temp_dir, f"{output_name}_{unique_id}")
        
        # Clean up transient sample docx
        try:
            if os.path.exists(filled_docx):
                os.remove(filled_docx)
        except Exception:
            pass

        return pdf_path
    except Exception as e:
        raise ValueError(f"Failed to render sample PDF preview: {str(e)}")
